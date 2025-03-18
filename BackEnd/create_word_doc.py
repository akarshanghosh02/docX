# -*- coding: utf-8 -*-
"""
Created on Thu Feb  3 12:30:56 2022
@author: bhushan.nehete
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import xml.etree.ElementTree as ET
import re
import string
from docx.shared import Pt  # Helps to specify font size
from docx.shared import RGBColor
import ComponentName
import readCSV
import create_bundle_sorter

# Retrieving extracted data from bundle sorter file
extQuery = create_bundle_sorter.getZoneValue

algorithmType = create_bundle_sorter.algoType

algorithmSoftParam = create_bundle_sorter.algorithmSoftParameter
algoTypeSoftparam = create_bundle_sorter.algorithmTypeSoftParameter
componentList = create_bundle_sorter.componentList
list_of_component = create_bundle_sorter.list_of_component
charType_list = create_bundle_sorter.charType
message_list = create_bundle_sorter.messages_list

dataArea = create_bundle_sorter.dataArea
uiMap = create_bundle_sorter.uiMap
batchControl = create_bundle_sorter.batchCntlComp
portalComp = create_bundle_sorter.portalComp
fkRefComp = create_bundle_sorter.FKRefComp
billSegTypComp = create_bundle_sorter.billSegTypComp
calcGrpComp = create_bundle_sorter.calcGrpComp
busFlgTypComp = create_bundle_sorter.busFlgTypComp
fieldComp = create_bundle_sorter.fieldComp
billMsgComp = create_bundle_sorter.billMsgComp
featCnfgComp = create_bundle_sorter.featCnfgComp
adjTypeComp = create_bundle_sorter.adjTypeComp
adjTypePrfComp = create_bundle_sorter.adjTypePrfComp
appSvcComp = create_bundle_sorter.appSvcComp
bktCnfgComp = create_bundle_sorter.bktCnfgComp
calcLineCatTypComp = create_bundle_sorter.calcLineCatTypComp
iwsSvcSOAPComp = create_bundle_sorter.iwsSvcSOAPComp
caseTypComp = create_bundle_sorter.caseTypComp
final_case_lfc = create_bundle_sorter.final_case_lfc
leadEvntTypComp = create_bundle_sorter.leadEvntTypComp
lookupComp = create_bundle_sorter.lookupComp
mainObjComp = create_bundle_sorter.mainObjComp
mgmtContentComp = create_bundle_sorter.mgmtContentComp
migrPlnComp = create_bundle_sorter.migrPlnComp
migrReqComp = create_bundle_sorter.migrReqComp
navKeyComp = create_bundle_sorter.navKeyComp
navOptComp = create_bundle_sorter.navOptComp
ntfTypComp = create_bundle_sorter.ntfTypComp
perCntTypComp = create_bundle_sorter.perCntTypComp
rateSchedComp = create_bundle_sorter.rateSchedComp
initiativeComp = create_bundle_sorter.initiativeComp
salesRepComp = create_bundle_sorter.salesRepComp
toDoTypComp = create_bundle_sorter.toDoTypComp
###############################################################
target_component=[]
zoneQuery=[]
algos = []
scriptData = {}
business_service = {}
business_object = {}
bo_Life_Cycle = {}
# Read CSV function call to get headings
df = readCSV.readCSVFiles()

def createDocument(bundle_data,selected_component):
    global target_component
    target_component = selected_component
    fetchData(bundle_data)

    doc = Document(df.loc['templateDocument'].value)
    query=zoneQuery
    # Iterating through paragraphs present in template
    for para in doc.paragraphs:
        # setRevisionControl(para)
        setSqlQuery(query, para, doc)
        setOUAFEntities(para)
        setExportableComp(para)
    doc.save(ComponentName.outputFileName)
    print("Document created")
def fetchData(bundle_data):
    entities = bundle_data["root"]["bundle"]["entities"]
    for entity in entities:
        componentType = entity['mo']
        component = entity['pk1']
        boData = entity['boData']
        if boData is not None:
            getBOData(boData, componentType, component)

# This is the main function where we iterate through template and create a document


def getBOData(boData, componentType, component):
    # boDataChildren = boData.childNodes
    global zoneQuery,algos,algorithmType,scriptData,business_object
    if componentType == 'CONTENT ZONE' and component in target_component:
        zoneQuery.append(create_bundle_sorter.getZoneValue(boData, component))
    if componentType == 'ALGORITHM' and component in target_component:
        algos.append(create_bundle_sorter.getAlgoSoftParam(boData, component))
    if componentType == 'ALG TYPE' and component in target_component:
        algorithmType.append(create_bundle_sorter.getAlgoTypeSoftParam(boData,component))
    if componentType == 'SCRIPT' and component in target_component:
        scriptData[component]=(create_bundle_sorter.getScriptData(boData,component))
    if componentType == 'F1-BUS OBJ' and component in target_component:
        business_object[component]=(create_bundle_sorter.getBusiness_Service_Object_Data(boData,component,componentType))



# This fuction create the Rivision Control table under Introduction
def setRevisionControl(para):
    # check the insertion point
    if 'Introduction' in para.text:
        para1 = insert_paragraph_after(para, "Revision Control")
        columneName = ['Cycle #', 'Version', 'Date', 'Author', 'JIRA Id', 'Details of change']
        table = create_table_header_and_row(row=1, column=6, style='Table Grid', list_of_headers=columneName)
        print('1')
        for i in range(2):  # change the code for revision control in the future
            row_cells = table.add_row().cells
            row_cells[0].text = "N/A"
            row_cells[3].text = "N/A"
            row_cells[5].text = "N/A"
        move_table_after(table, para1)


# This fuction tends to put all OUAF Entities present in the bundle
def setOUAFEntities(para):
    global algos, algorithmType
    # check the insertion point and call the functions which ones presents in bundle
    if 'OUAF Entities' in para.text:
        if algos:
            setAlgo(para)
        if algorithmType:
            setAlgoType(para)
        if scriptData:
            setScriptEntity(para)
        if business_object:
            setBsBoSchema(para, business_object, column_header=ComponentName.businessObject)
        if business_service:
            setBsBoSchema(para, business_service, column_header=ComponentName.businessService)
        if charType_list:
            setCharType(para)
        if message_list:
            setMessage(para)
        if dataArea:
            setDataArea(para)
        # if uiMap:
        #     setUiMap(para)
        if batchControl:
            setBatchCntl(para)
        if portalComp:
            setPortal(para)
        if fkRefComp:
            setFKRef(para)
        if billSegTypComp:
            setBillSegTyp(para)
        if calcGrpComp:
            setCalcGrp(para)
        if busFlgTypComp:
            setBusFlgTyp(para)
        if fieldComp:
            setField(para)
        if billMsgComp:
            setBillMsg(para)
        if featCnfgComp:
            setFeatCnfg(para)
        if adjTypePrfComp:
            setAdjTypePrf(para)
        if adjTypeComp:
            setAdjType(para)
        if appSvcComp:
            setAppService(para)
        if bktCnfgComp:
            setBktCnfg(para)
        if calcLineCatTypComp:
            setCalcLineCatTyp(para)
        if iwsSvcSOAPComp:
            setIwsSvcSOAP(para)
        if caseTypComp:
            setCaseTyp(para)
        # if leadEvntTypComp:
        #     setLeadEvntTyp(para)
        if lookupComp:
            setLookup(para)
        if mainObjComp:
            setMainObj(para)
        if mgmtContentComp:
            setMgmtContent(para)
        if migrPlnComp:
            setMigrPln(para)
        if migrReqComp:
            setMigrReq(para)
        if navKeyComp:
            setNavKey(para)
        if navOptComp:
            setNavOpt(para)
        if ntfTypComp:
            setNtfTyp(para)
        if perCntTypComp:
            setPerCntTyp(para)
        if rateSchedComp:
            setRateSched(para)
        if initiativeComp:
            setInitiative(para)
        if salesRepComp:
            setSalesRep(para)
        if toDoTypComp:
            setToDoTyp(para)


def setToDoTyp(para):
    # Creates table header
    p1 = tableHeader(para, "To Do Type")

    for toDoTyp in toDoTypComp:
        # Create  table  with component name and properties
        table = create_table_header(toDoTyp, ["Property", "Value"])
        # Merge first main header row with no. of sub-header count
        row = table.row_cells(0)
        row[0].merge(row[1])

        # Get common component fields and add to table rows
        value = toDoTypComp[toDoTyp][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        # Bind table to component header
        move_table_after(table, p1)

        # Create  table  with main and sub headers
        table1 = create_table_header("To Do Role", ["Role", "Default"])
        # Merge first main header row with no. of sub-header count
        row = table1.row_cells(0)
        row[0].merge(row[1])

        value = toDoTypComp[toDoTyp][1]
        values = value.get('ToDoTypRole')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('toDoRole') if entity.get('toDoRole') else ""
                row_cells[1].text = entity.get('shouldUseAsDefault') if entity.get('shouldUseAsDefault') else ""
        else:
            table1.add_row().cells
        # Merge tables
        move_after_table(table1, table)

        # Create  table  with main and sub headers
        table2 = create_table_header("Sort Keys",
                                     ["Sequence", "Description", "Default", "Order"])
        # Merge first main header row with no. of sub-header count
        row = table2.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3])

        # Get common component fields and add to table rows
        value = toDoTypComp[toDoTyp][1]
        values = value.get('ToDoTypSortKey')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('description') if entity.get('description') else ""
                row_cells[2].text = entity.get('shouldUseAsDefault') if entity.get('shouldUseAsDefault') else ""
                row_cells[3].text = entity.get('order') if entity.get('order') else ""
        else:
            table2.add_row().cells
        # Merge tables
        move_after_table(table2, table1)

        # Create  table  with main and sub headers
        table3 = create_table_header("Drill Keys",
                                     ["Sequence", "Table", "Field"])
        # Merge first main header row with no. of sub-header count
        row = table3.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        # Get common component fields and add to table rows
        value = toDoTypComp[toDoTyp][1]
        values = value.get('ToDoTypDrillKey')
        if values:
            for entity in values:
                row_cells = table3.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('table') if entity.get('table') else ""
                row_cells[2].text = entity.get('field') if entity.get('field') else ""
        else:
            table3.add_row().cells
        # Merge tables
        move_after_table(table3, table2)

        # Create  table  with main and sub headers
        table4 = create_table_header("Characteristic Control",
                                     ["Sequence", "Characteristic Type", "is Required",
                                      "Default", "Characteristic Value"])
        # Merge first main header row with no. of sub-header count
        row = table4.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])

        # Get common component fields and add to table rows
        value = toDoTypComp[toDoTyp][1]
        values = value.get('ToDoTypCharCntl')
        if values:
            for entity in values:
                row_cells = table4.add_row().cells
                row_cells[0].text = entity.get('sortSequence') if entity.get('sortSequence') else ""
                row_cells[1].text = entity.get('characteristicType') if entity.get('characteristicType') else ""
                row_cells[2].text = entity.get('isRequired') if entity.get('isRequired') else ""
                row_cells[3].text = entity.get('shouldUseAsDefault') if entity.get('shouldUseAsDefault') else ""
                row_cells[4].text = entity.get('characteristicValue') if entity.get('characteristicValue') else ""
        else:
            table4.add_row().cells
        # Merge tables
        move_after_table(table4, table3)

        # Create  table  with main and sub headers
        table5 = create_table_header("Characteristics",
                                     ["Sequence", "Characteristic Type", "Characteristic Value"])
        # Merge first main header row with no. of sub-header count
        row = table5.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        # Get common component fields and add to table rows
        value = toDoTypComp[toDoTyp][1]
        values = value.get('ToDoTypChar')
        if values:
            for entity in values:
                row_cells = table5.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('characteristicType') if entity.get('characteristicType') else ""
                row_cells[2].text = entity.get('characteristicValue') if entity.get('characteristicValue') else ""
        else:
            table5.add_row().cells
        # Merge tables
        move_after_table(table5, table4)
        # Add space after table
        spaceAfterTable(table5)


def setSalesRep(para):
    # Creates table header
    p1 = tableHeader(para, "Sales Representation")
    for salesRep in salesRepComp:
        # Create  table  with main and sub headers
        table = create_table_header(salesRep, ["Property", "Value"])
        # columnWidth(table, 0, 190200)
        # Merge component name header row
        row = table.row_cells(0)
        row[0].merge(row[1])

        # Get common component fields and add to table rows
        value = salesRepComp[salesRep][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        # Add space after table
        spaceAfterTable(table)


def setInitiative(para):
    # Creates table header
    p1 = tableHeader(para, "Initiative")

    for initiative in initiativeComp:
        # Create  table  with main and sub headers
        table = create_table_header(initiative, ["Property", "Value"])
        # Merge component name header row
        row = table.row_cells(0)
        row[0].merge(row[1])

        # Get common component fields and add to table rows
        value = initiativeComp[initiative][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        # Bind table to component header
        move_table_after(table, p1)

        # Create  table  with main and sub headers
        table1 = create_table_header("Lead Definition",
                                     ["Sequence", "Event Type", "Days After Lead Creation"])
        # Merge first main header row with no. of sub-header count
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        # Get extra component fields identified by name given at time of extracting fields
        value = initiativeComp[initiative][1]
        values = value.get('initiativeLeadDef')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('eventSequence') if entity.get('eventSequence') else ""
                row_cells[1].text = entity.get('leadEventType') if entity.get('leadEventType') else ""
                row_cells[2].text = entity.get('daysAfterLeadCreation') if entity.get('daysAfterLeadCreation') else ""
        else:
            table1.add_row().cells
        # Merge tables
        move_after_table(table1, table)
        # Add space after table
        spaceAfterTable(table1)


def setRateSched(para):
    p1 = tableHeader(para, "Rate Schedule")

    for rateSched in rateSchedComp:
        table = create_table_header(rateSched, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = rateSchedComp[rateSched][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Post-Processing",
                                     ["Sequence", "Description", "Calculation Group"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = rateSchedComp[rateSched][1]
        values = value.get('rtSchPostPrc')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('description') if entity.get('description') else ""
                row_cells[2].text = entity.get('calculationGroup') if entity.get('calculationGroup') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Meter Configuration Type")

        value = rateSchedComp[rateSched][1]
        values = value.get('rtSchMeterCnfgTyp')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('meterConfigurationType') if entity.get('meterConfigurationType') else ""
        else:
            table2.add_row().cells
        move_after_table(table2, table1)

        table3 = create_table_header("Rate Version",
                                     ["Effective Date", "Description", "Calculation Group"])
        row = table3.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = rateSchedComp[rateSched][1]
        values = value.get('rtSchRtVer')
        if values:
            for entity in values:
                row_cells = table3.add_row().cells
                row_cells[0].text = entity.get('effectiveDate') if entity.get('effectiveDate') else ""
                row_cells[1].text = entity.get('rateVersionDescription') if entity.get('rateVersionDescription') else ""
                row_cells[2].text = entity.get('calculationGroup') if entity.get('calculationGroup') else ""
        else:
            table3.add_row().cells
        move_after_table(table3, table2)

        table4 = create_table_header("Pre-Processing",
                                     ["Sequence", "Description", "Calculation Group"])
        row = table4.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = rateSchedComp[rateSched][1]
        values = value.get('rtSch')
        if values:
            for entity in values:
                row_cells = table4.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('description') if entity.get('description') else ""
                row_cells[2].text = entity.get('calculationGroup') if entity.get('calculationGroup') else ""
        else:
            table4.add_row().cells
        move_after_table(table4, table3)

        table5 = create_table_header("Messages",
                                     ["Bill Message", "Start Date", "End Date"])
        row = table5.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = rateSchedComp[rateSched][1]
        values = value.get('rtSchBillMsg')
        if values:
            for entity in values:
                row_cells = table5.add_row().cells
                row_cells[0].text = entity.get('billMessage') if entity.get('billMessage') else ""
                row_cells[1].text = entity.get('startDate') if entity.get('startDate') else ""
                row_cells[2].text = entity.get('endDate') if entity.get('endDate') else ""
        else:
            table5.add_row().cells
        move_after_table(table5, table4)
        spaceAfterTable(table5)


def setPerCntTyp(para):
    p1 = tableHeader(para, "Person Contact Type")

    for perCntTyp in perCntTypComp:
        table = create_table_header(perCntTyp, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = perCntTypComp[perCntTyp][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Delivery Types")

        value = perCntTypComp[perCntTyp][1]
        values = value.get('perCntTypDelvTyp')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('deliveryType') if entity.get('deliveryType') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Characteristics")

        value = perCntTypComp[perCntTyp][1]
        values = value.get('perCntTypChar')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('characteristicValueForeignKey1') if entity.get(
                    'characteristicValueForeignKey1') else ""
        else:
            table2.add_row().cells
        move_after_table(table2, table1)
        spaceAfterTable(table2)


def setNtfTyp(para):
    p1 = tableHeader(para, "Notification Type")

    for ntfTyp in ntfTypComp:
        table = create_table_header(ntfTyp, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = ntfTypComp[ntfTyp][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Algorithms",
                                     ["Sequence", "System Event", "Algorithm"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = ntfTypComp[ntfTyp][1]
        values = value.get('ntfTypAlgo')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('systemEvent') if entity.get('systemEvent') else ""
                row_cells[2].text = entity.get('algorithm') if entity.get('algorithm') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Person Contacts")

        value = ntfTypComp[ntfTyp][1]
        values = value.get('ntfTypPer')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('personContactType') if entity.get('personContactType') else ""
        else:
            table2.add_row().cells
        move_after_table(table2, table1)
        spaceAfterTable(table2)


def setNavOpt(para):
    p1 = tableHeader(para, "Migration Request")

    for navOpt in navOptComp:
        table = create_table_header(navOpt, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = navOptComp[navOpt][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Option Usage")
        value = navOptComp[navOpt][1]
        values = value.get('navOptUsage')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('navigationOptionUsage') if entity.get('navigationOptionUsage') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Option Context", ["Field", "is Key Field"])
        row = table2.row_cells(0)
        row[0].merge(row[1])

        value = navOptComp[navOpt][1]
        values = value.get('navOptCntxt')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('field') if entity.get('field') else ""
                row_cells[1].text = entity.get('isKeyField') if entity.get('isKeyField') else ""
        else:
            table2.add_row().cells
        move_after_table(table2, table1)
        spaceAfterTable(table2)


def setNavKey(para):
    p1 = tableHeader(para, "Navigation Key")

    for navKey in navKeyComp:
        table = create_table_header(navKey, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = navKeyComp[navKey][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)
        spaceAfterTable(table)


def setMigrReq(para):
    p1 = tableHeader(para, "Migration Request")

    for migrReq in migrReqComp:
        table = create_table_header(migrReq, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = migrReqComp[migrReq][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Instructions",
                                     ["Migration Plan", "Selection Type", "SQL Statement Text", "Instructions Entity"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3])

        value = migrReqComp[migrReq][1]
        values = value.get('migrReqInstruc')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('migrationPlan') if entity.get('migrationPlan') else ""
                row_cells[1].text = entity.get('selectionType') if entity.get('selectionType') else ""
                row_cells[2].text = entity.get('sqlStatementText') if entity.get('sqlStatementText') else ""
                row_cells[3].text = entity.get('primaryKeyValue1') if entity.get('primaryKeyValue1') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)


def setMigrPln(para):
    p1 = tableHeader(para, "Migration Plan")

    for migrPln in migrPlnComp:
        table = create_table_header(migrPln, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = migrPlnComp[migrPln][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)

        # table1 = create_table_header("Instructions",
        #                              ["Sequence", "instructionType", "parentInstructionSequence", "businessObject",
        #                               "nextMigrationPlan", "traversalCriteriaType", "constraintId",
        #                               "referringConstraintOwner", "description", "sqlTraversalCriteria",
        #                               "xPathTraversalCriteria"])
        # row = table1.row_cells(0)
        # row[0].merge(row[1]).merge(row[2])
        #
        # # value = migrPlnComp[migrPln][1]
        # # values = value.get('migrPlnInstruc')
        # # if values:
        # #     for entity in values:
        # #         row_cells = table1.add_row().cells
        # #         row_cells[0].text = entity.get('Sequence') if entity.get('Sequence') else ""
        # #         row_cells[1].text = entity.get('instructionType')
        # #         row_cells[2].text = entity.get('parentInstructionSequence') if entity.get('parentInstructionSequence') else ""
        # #         row_cells[3].text = entity.get('businessObject')
        # #         row_cells[4].text = entity.get('nextMigrationPlan') if entity.get('nextMigrationPlan') else ""
        # #         row_cells[5].text = entity.get('traversalCriteriaType') if entity.get('traversalCriteriaType') else ""
        # #         row_cells[6].text = entity.get('constraintId') if entity.get('constraintId') else ""
        # #         row_cells[7].text = entity.get('referringConstraintOwner') if entity.get('referringConstraintOwner') else ""
        # #         row_cells[8].text = entity.get('description') if entity.get('description') else ""
        # #         if entity.get('sqlTraversalCriteria'):
        # #             row_cells[9].text = entity.get('sqlTraversalCriteria') if entity.get('sqlTraversalCriteria') else ""
        # #         if entity.get('xPathTraversalCriteria'):
        # #             row_cells[10].text = entity.get('xPathTraversalCriteria') if entity.get('xPathTraversalCriteria') else ""
        # # else:
        # #     table1.add_row().cells
        # # move_after_table(table1, table)
        spaceAfterTable(table)


def setMgmtContent(para):
    p1 = tableHeader(para, "Management Content")

    for mgmtContent in mgmtContentComp:
        table = create_table_header(mgmtContent, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = mgmtContentComp[mgmtContent][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setMainObj(para):
    p1 = tableHeader(para, "Maintenance Object")

    for mainObj in mainObjComp:
        table = create_table_header(mainObj, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = mainObjComp[mainObj][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Maintenance Object Table",
                                     ["Table", "Table Role", "Parent Constraint Id"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = mainObjComp[mainObj][1]
        values = value.get('mainObjTbl')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('table') if entity.get('table') else ""
                row_cells[1].text = entity.get('tableRole') if entity.get('tableRole') else ""
                row_cells[2].text = entity.get('parentConstraintId') if entity.get('parentConstraintId') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Algorithms",
                                     ["Sequence", "Event", "Algorithm"])
        row = table2.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = mainObjComp[mainObj][1]
        values = value.get('mainObjAlgo')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('event') if entity.get('event') else ""
                row_cells[2].text = entity.get('algorithm') if entity.get('algorithm') else ""
        else:
            table2.add_row().cells
        move_after_table(table2, table1)

        table3 = create_table_header("Options",
                                     ["Option Type", "sequence", "Option Value"])
        row = table3.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = mainObjComp[mainObj][1]
        values = value.get('mainObjOpt')
        if values:
            for entity in values:
                row_cells = table3.add_row().cells
                row_cells[0].text = entity.get('maintenanceObjectOptionType') if entity.get(
                    'maintenanceObjectOptionType') else ""
                row_cells[1].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[2].text = entity.get('maintenanceObjectOptionValue') if entity.get(
                    'maintenanceObjectOptionValue') else ""
        else:
            table3.add_row().cells
        move_after_table(table3, table2)
        spaceAfterTable(table3)


def setLookup(para):
    p1 = tableHeader(para, "Lookup")

    for lookup in lookupComp:
        table = create_table_header(lookup, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = lookupComp[lookup][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Lookup Value",
                                     ["Field Value", "Description", "Value Name"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = lookupComp[lookup][1]
        values = value.get('lookVal')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('fieldValue') if entity.get('fieldValue') else ""
                row_cells[1].text = entity.get('description') if entity.get('description') else ""
                row_cells[2].text = entity.get('valueName') if entity.get('valueName') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)


def setLeadEvntTyp(para):
    p1 = tableHeader(para, "Lead Event Type")

    for leadEvntTyp in leadEvntTypComp:
        table = create_table_header(leadEvntTyp, ["Property", "Value"])
        # columnWidth(table, 0, 690200)
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = leadEvntTypComp[leadEvntTyp][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setCaseTyp(para):
    p1 = tableHeader(para, "Case Type")

    for caseTyp in caseTypComp:
        table = create_table_header(caseTyp, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = caseTypComp[caseTyp][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Algorithm",
                                     ["Sequence", "System Event", "Algorithm"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = caseTypComp[caseTyp][1]
        values = value.get('caseTypAlgo')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('caseTypeSystemEvent') if entity.get('caseTypeSystemEvent') else ""
                row_cells[2].text = entity.get('algorithm') if entity.get('algorithm') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Characteristics",
                                     ["Sequence", "Characteristic Type", "is Required", "Default",
                                      "Characteristic Value"])
        row = table2.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])

        value = caseTypComp[caseTyp][1]
        values = value.get('caseTypChar')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('sortSequence') if entity.get('sortSequence') else ""
                row_cells[1].text = entity.get('characteristicType') if entity.get('characteristicType') else ""
                row_cells[2].text = entity.get('isRequired') if entity.get('isRequired') else ""
                row_cells[3].text = entity.get('shouldUseAsDefault') if entity.get('shouldUseAsDefault') else ""
                row_cells[4].text = entity.get('characteristicValue') if entity.get('characteristicValue') else ""
        else:
            table1.add_row().cells
        move_after_table(table2, table1)

        if final_case_lfc:
            tbl = table2._tbl
            p1 = tableHeader(para, "Case Type Life Cycle")
            tbl.addnext(p1._p)
            setCaseLfc(p1)


def setCaseLfc(para):
    for cLfc in reversed(final_case_lfc):
        table = create_table_header(cLfc, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = final_case_lfc[cLfc][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, para)

        table1 = create_table_header("Algorithm",
                                     ["Sequence", "System Event", "Algorithm"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = final_case_lfc[cLfc][1]
        values = value.get('caseLfcAlgo')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('caseStatusSystemEvent') if entity.get('caseStatusSystemEvent') else ""
                row_cells[2].text = entity.get('algorithm') if entity.get('algorithm') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Characteristic Types")

        value = final_case_lfc[cLfc][1]
        values = value.get('caseLfcCharRule')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('characteristicType') if entity.get('characteristicType') else ""
        else:
            table2.add_row().cells
        move_after_table(table2, table1)

        table3 = create_table_header("Next Status",
                                     ["Status", "Transition Role", "Sequence", "Default", "Transition Condition",
                                      "Description"])
        row = table3.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4]).merge(row[5])

        value = final_case_lfc[cLfc][1]
        values = value.get('nxtState')
        if values:
            for entity in values:
                row_cells = table3.add_row().cells
                row_cells[0].text = entity.get('nextStatus') if entity.get('nextStatus') else ""
                row_cells[1].text = entity.get('transitionRole') if entity.get('transitionRole') else ""
                row_cells[2].text = entity.get('sortSequence') if entity.get('sortSequence') else ""
                row_cells[3].text = entity.get('shouldUseAsDefault') if entity.get('shouldUseAsDefault') else ""
                row_cells[4].text = entity.get('transitionCondition') if entity.get('transitionCondition') else ""
                row_cells[5].text = entity.get('nextStatus2') if entity.get('nextStatus2') else ""
        else:
            table3.add_row().cells
        move_after_table(table3, table2)
        spaceAfterTable(table3)


def setIwsSvcSOAP(para):
    p1 = tableHeader(para, "SOAP Inbound Web Services")

    for iwsSvcSOAP in iwsSvcSOAPComp:
        table = create_table_header(iwsSvcSOAP, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = iwsSvcSOAPComp[iwsSvcSOAP][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Operations",
                                     ["Operation Name", "Schema Name", "Schema Type", "Transaction Type"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3])

        value = iwsSvcSOAPComp[iwsSvcSOAP][1]
        values = value.get('iwsSvcSOAPOpt')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('operationName') if entity.get('operationName') else ""
                row_cells[1].text = entity.get('schemaName') if entity.get('schemaName') else ""
                row_cells[2].text = entity.get('schemaType') if entity.get('schemaType') else ""
                row_cells[3].text = entity.get('transactionType') if entity.get('transactionType') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)


def setCalcLineCatTyp(para):
    p1 = tableHeader(para, "Calculation Line Category")

    for calcLineCatTyp in calcLineCatTypComp:
        table = create_table_header(calcLineCatTyp, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = calcLineCatTypComp[calcLineCatTyp][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Calculation Line Category Value", ["Category Values", "description"])
        row = table1.row_cells(0)
        row[0].merge(row[1])

        value = calcLineCatTypComp[calcLineCatTyp][1]
        values = value.get('calcLineCatValue')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('calcLineCategoryValue') if entity.get('calcLineCategoryValue') else ""
                row_cells[1].text = entity.get('description') if entity.get('description') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)


def setBktCnfg(para):
    p1 = tableHeader(para, "Bucket Configuration")

    for bktCnfg in bktCnfgComp:
        table = create_table_header(bktCnfg, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = bktCnfgComp[bktCnfg][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Bucket Configuration Value",
                                     ["Sequence", "Start Range", "End Range", "Description"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3])

        value = bktCnfgComp[bktCnfg][1]
        values = value.get('bktCnfgValue')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('startRange') if entity.get('startRange') else ""
                row_cells[2].text = entity.get('endRange') if entity.get('endRange') else ""
                row_cells[3].text = entity.get('description') if entity.get('description') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)


def setAppService(para):
    p1 = tableHeader(para, "Application Service")

    for appSvc in appSvcComp:
        table = create_table_header(appSvc, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = appSvcComp[appSvc][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Access Modes")

        value = appSvcComp[appSvc][1]
        values = value.get('appSvcMode')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('accessMode') if entity.get('accessMode') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)


def setAdjType(para):
    # p1 = insert_paragraph_after(para, None, None)
    # run = p1.add_run("\n" + "Adjustment Type" + ComponentName.seperator)
    # run.bold = True
    # run.font.size = Pt(16)
    # run.font.color.rgb = RGBColor(197, 125, 90)
    # run.font.name = 'PT Sans Narrow'
    p1 = tableHeader(para, "Adjustment Type")

    for adjType in adjTypeComp:
        table = create_table_header(adjType, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = adjTypeComp[adjType][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Characteristics", ["Characteristic Type", "Characteristic Value"])
        row = table1.row_cells(0)
        row[0].merge(row[1])

        value = adjTypeComp[adjType][1]
        values = value.get('adjTypeChar')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('characteristicType') if entity.get('characteristicType') else ""
                row_cells[1].text = entity.get('characteristicValue') if entity.get('characteristicValue') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Algorithm", ["Sequence", "Algorithm Entity", "Algorithm"])
        row = table2.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = adjTypeComp[adjType][1]
        values = value.get('adjTypeAlgo')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('adjustmentTypeAlgorithmEntity') if entity.get(
                    'adjustmentTypeAlgorithmEntity') else ""
                row_cells[2].text = entity.get('algorithm') if entity.get('algorithm') else ""
        else:
            table2.add_row().cells
        move_after_table(table2, table1)

        table3 = create_table_header("Adjustment Characteristic",
                                     ["Sequence", "Characteristic Type", "Is Required", "Default",
                                      "Characteristic Value"])
        row = table3.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])

        value = adjTypeComp[adjType][1]
        values = value.get('adjTypeTempChar')
        if values:
            for entity in values:
                row_cells = table3.add_row().cells
                row_cells[0].text = entity.get('sortSequence') if entity.get('sortSequence') else ""
                row_cells[1].text = entity.get('characteristicType') if entity.get('characteristicType') else ""
                row_cells[2].text = entity.get('isRequired') if entity.get('isRequired') else ""
                row_cells[3].text = entity.get('shouldUseAsDefault') if entity.get('shouldUseAsDefault') else ""
                row_cells[4].text = entity.get('characteristicValue') if entity.get('characteristicValue') else ""
        else:
            table3.add_row().cells
        move_after_table(table3, table2)
        spaceAfterTable(table3)


def setAdjTypePrf(para):
    p1 = tableHeader(para, "Adjustment Type Profile")

    for adjTypePrf in adjTypePrfComp:
        table = create_table_header(adjTypePrf, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = adjTypePrfComp[adjTypePrf][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Profiles")
        value = adjTypePrfComp[adjTypePrf][1]
        values = value.get('adjTypePrf')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('adjustmentType') if entity.get('adjustmentType') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)


def setFKRef(para):
    p1 = tableHeader(para, "Foreign Key Reference")

    for fkref in fkRefComp:
        table = create_table_header(fkref, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = fkRefComp[fkref][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setBillSegTyp(para):
    p1 = tableHeader(para, "Bill Segment Type")

    for billSegTyp in billSegTypComp:
        table = create_table_header(billSegTyp, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = billSegTypComp[billSegTyp][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setCalcGrp(para):
    p1 = tableHeader(para, "Calculation Group")

    for calcGrp in calcGrpComp:
        table = create_table_header(calcGrp, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = calcGrpComp[calcGrp][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setBusFlgTyp(para):
    p1 = tableHeader(para, "Business Flag Type")

    for busFlgTyp in busFlgTypComp:
        table = create_table_header(busFlgTyp, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = busFlgTypComp[busFlgTyp][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setField(para):
    p1 = tableHeader(para, "Field")

    for field in fieldComp:
        table = create_table_header(field, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = fieldComp[field][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setBillMsg(para):
    p1 = tableHeader(para, "Bill Message")

    for billMsg in billMsgComp:
        table = create_table_header(billMsg, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = billMsgComp[billMsg][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setFeatCnfg(para):
    p1 = tableHeader(para, "Feature Configuration")

    for featCnfg in featCnfgComp:
        table = create_table_header(featCnfg, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = featCnfgComp[featCnfg][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)

        table1 = create_table_header("Options", ["Sequence", "Option Type", "Value"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = featCnfgComp[featCnfg][1]
        values = value.get('featCnfgOpt')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('optionType') if entity.get('optionType') else ""
                row_cells[2].text = entity.get('value') if entity.get('value') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Messages",
                                     ["Feature Message Category", "Message", "Message Category", "Message Number"])
        row = table2.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3])

        value = featCnfgComp[featCnfg][1]
        values = value.get('featCnfgMsg')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('workforceManagementMessageCategory')
                row_cells[1].text = entity.get('wfmMessage')
                row_cells[2].text = entity.get('messageCategory')
                row_cells[3].text = entity.get('messageNumber')
        else:
            table2.add_row().cells
        move_after_table(table2, table1)
        spaceAfterTable(table2)


def setPortal(para):
    p1 = tableHeader(para, "Portal")

    for portal in portalComp:
        table = create_table_header(portal, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = portalComp[portal][0]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)

        table1 = create_table_header("Zones", ['Sort Sequence', 'Zone', 'Can Display'])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = portalComp[portal][1]
        values = value.get('zone')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('sortSequence') if entity.get('sortSequence') else ""
                row_cells[1].text = entity.get('zone') if entity.get('zone') else ""
                row_cells[2].text = entity.get('canDisplay') if entity.get('canDisplay') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)

        table2 = create_table_header("Options",
                                     ['Sequence', 'Portal Option Type', 'Portal Option Value'])
        row = table2.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = portalComp[portal][1]
        values = value.get('option')
        if values:
            for entity in values:
                row_cells = table2.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('portalOptionType') if entity.get('portalOptionType') else ""
                row_cells[2].text = entity.get('portalOptionValue') if entity.get('portalOptionValue') else ""
        else:
            table2.add_row().cells
        move_after_table(table2, table1)
        spaceAfterTable(table2)


# This block of code uses only one table to print the details where other functions merge multiple tables
def setBatchCntl(para):
    doc = Document(df.loc['templateDocument'].value)
    p1 = tableHeader(para, "Batch Control")
    for batch in batchControl:
        table = doc.add_table(rows=2, cols=5, style='Table Grid')
        hdr_cells = table.rows[0].cells
        count = rowcount(table)
        hdr_cells[0].paragraphs[0].add_run(batch).bold = True
        row = table.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])
        backgroundColor(table=table, row=0, column=0)

        table_headers = ["Property", "", "", "", "Value"]
        c = len(table_headers)
        count = rowcount(table)
        for i in range(c):
            hdr_cells = table.rows[1].cells
            hdr_cells[i].paragraphs[0].add_run(table_headers[i]).bold = True
            row = table.row_cells(1)
            row[0].merge(row[1])
            row[2].merge(row[3]).merge(row[4])
            backgroundColor(table=table, row=1, column=i)

        value = batchControl[batch][0]
        for entity in value:
            row_cells = table.add_row().cells
            row = row_cells
            row[0].merge(row[1])
            row[2].merge(row[3]).merge(row[4])
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[2].text = value[entity]

        value = batchControl[batch][1]
        count = rowcount(table)
        if 'Parameters' in value:
            table.add_row().cells
            hdr_cells = table.rows[count].cells
            run = hdr_cells[0].paragraphs[0].add_run('Parameters')
            run.bold = True
            row = table.row_cells(count)
            row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])
            backgroundColor(table=table, row=count, column=0)

            table_headers = ['Sequence', 'Batch Parameter Name', 'Batch Parameter Value', 'is Required', 'Description']
            c = len(table_headers)
            count = rowcount(table)
            table.add_row().cells
            for i in range(c):
                hdr_cells = table.rows[count].cells
                run = hdr_cells[i].paragraphs[0].add_run(table_headers[i])
                run.bold = True
            for i in range(c):
                backgroundColor(table=table, row=count, column=i)

            parameters = value.get('Parameters')
            for para in parameters:
                row_cells = table.add_row().cells
                row_cells[0].text = para.get('sequence')
                row_cells[1].text = para.get('batchParameterName')
                if para.get('batchParameterValue'):
                    row_cells[2].text = para.get('batchParameterValue')
                row_cells[3].text = para.get('isRequired')
                row_cells[4].text = para.get('description')

        count = rowcount(table)
        if 'Algorithm' in value:
            table.add_row().cells
            hdr_cells = table.rows[count].cells
            run = hdr_cells[0].paragraphs[0].add_run('Algorithms')
            run.bold = True
            row = table.row_cells(count)
            row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])
            backgroundColor(table=table, row=count, column=0)

            table_headers = ['Batch Control System Event', "", 'Sequence', "", 'Algorithm']
            c = len(table_headers)
            count = rowcount(table)
            table.add_row().cells
            for i in range(c):
                hdr_cells = table.rows[count].cells
                row = hdr_cells
                row[0].merge(row[1])
                row[3].merge(row[4])
                run = hdr_cells[i].paragraphs[0].add_run(table_headers[i])
                run.bold = True
            for i in range(c):
                backgroundColor(table=table, row=count, column=i)
            algorithm = value.get('Algorithm')
            for algo in algorithm:
                row_cells = table.add_row().cells
                row = row_cells
                row[0].merge(row[1])
                row[3].merge(row[4])
                row_cells[0].text = algo.get('batchControlSystemEvent')
                row_cells[2].text = algo.get('sequence')
                row_cells[3].text = algo.get('algorithm')

        move_table_after(table, p1)
        spaceAfterTable(table)


def setBsBoSchema(para, bs_bo_dict, column_header):
    header = None
    scm = None
    if column_header == ComponentName.businessService:
        header = ComponentName.businessService
    elif column_header == ComponentName.businessObject:
        header = ComponentName.businessObject

    p1 = tableHeader(para, header)
    table_headers = ["Property", "Value"]
    c = len(table_headers)
    table = create_table_header_and_row(row=1, column=c, style='Table Grid', list_of_headers=table_headers)
    for i in range(c):
        backgroundColor(table=table, row=0, column=i)
    for bs_bo in bs_bo_dict:
        value = bs_bo_dict[bs_bo]
        for entity in value:
            # businessService = value[entity]
            if entity in ['description','applicationServiceId','maintenanceObject','longDescription']:

                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
            elif entity == 'schemaDefinition':
                scm = value[entity]
    move_table_after(table, p1)

    if column_header == ComponentName.businessObject:
        table2 = create_table_header("Algorithms", ["System Event", "Seq", "Algorithm"])
        row = table2.row_cells(0)
        row[0].merge(row[1]).merge(row[2])
        for bs_bo in bs_bo_dict:
            value = bs_bo_dict[bs_bo]['businessObjectAlgorithm']
            for res in value:
                table2.add_row().cells
                row_cells[0].text = res['event']
                row_cells[1].text = res['sequence']
                #row_cells[2].text= res['algorithm']

        move_after_table(table2, table)

    table_headers = ["Schema"]
    c = len(table_headers)
    table3 = create_table_header_and_row(row=1, column=c, style='Table Grid', list_of_headers=table_headers)
    row_cells = table3.add_row().cells
    # row_cells[0].text = xmlPrettyPrint(scm)
    row_cells[0].text = scm

    backgroundColor(table3, 0, 0)

    # row_cells[0].text = scm
    if column_header == ComponentName.businessObject:
        for i in range(c):
            backgroundColor(table=table2, row=0, column=i)
        move_after_table(table3, table2)
    else:
        move_after_table(table3, table)

    if header == ComponentName.businessObject:
        tbl = table3._tbl
        p1 = tableHeader(para, "Business Object Lifecycle")
        tbl.addnext(p1._p)
        setBoLifeCycle(p1)


# This block of code uses only one table to print the details where other functions merge multiple tables
def setBoLifeCycle(para):
    doc = Document(df.loc['templateDocument'].value)
    for bolfc in reversed(bo_Life_Cycle):
        table = doc.add_table(rows=2, cols=6, style='Table Grid')
        hdr_cells = table.rows[0].cells
        run = hdr_cells[0].paragraphs[0].add_run(bolfc)
        run.bold = True
        row = table.row_cells(0)
        row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4]).merge(row[5])
        backgroundColor(table=table, row=0, column=1)

        table_headers = ["Property", "", "Value"]
        c = len(table_headers)
        for i in range(c):
            hdr_cells = table.rows[1].cells
            row = hdr_cells
            row[0].merge(row[1])
            row[2].merge(row[3]).merge(row[4]).merge(row[5])
            run = hdr_cells[i].paragraphs[0].add_run(table_headers[i])
            run.bold = True
            backgroundColor(table=table, row=1, column=0)
            backgroundColor(table=table, row=1, column=2)

        value = bo_Life_Cycle[bolfc][0]
        for entity in value:
            row_cells = table.add_row().cells
            row = row_cells
            row[0].merge(row[1])
            row[2].merge(row[3]).merge(row[4]).merge(row[5])
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[2].text = value[entity]

        value = bo_Life_Cycle[bolfc][1]
        count = rowcount(table)
        if 'Algorithm' in value:
            table.add_row().cells
            hdr_cells = table.rows[count].cells
            run = hdr_cells[0].paragraphs[0].add_run('Algorithms')
            run.bold = True
            row = table.row_cells(count)
            row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4]).merge(row[5])
            backgroundColor(table=table, row=count, column=0)

            table_headers = ["System Event", "", "Seq", "", "", "Algorithm"]
            c = len(table_headers)
            count = rowcount(table)
            table.add_row().cells
            for i in range(c):
                hdr_cells = table.rows[count].cells
                row = hdr_cells
                row[0].merge(row[1])
                row[3].merge(row[4]).merge(row[5])
                run = hdr_cells[i].paragraphs[0].add_run(table_headers[i])
                run.bold = True
            for i in range(c):
                backgroundColor(table=table, row=count, column=i)

            algo = value.get('Algorithm')
            for al in algo:
                row_cells = table.add_row().cells
                row = row_cells
                row[0].merge(row[1])
                row[3].merge(row[4]).merge(row[5])
                row_cells[0].text = al.get('event')
                row_cells[2].text = al.get('sequence')
                row_cells[3].text = al.get('algorithm')

        count = rowcount(table)
        if 'nxtstatus' in value:
            table.add_row().cells
            hdr_cells = table.rows[count].cells
            run = hdr_cells[0].paragraphs[0].add_run('Next Statuses')
            run.bold = True
            row = table.row_cells(count)
            row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4]).merge(row[5])
            backgroundColor(table=table, row=count, column=0)

            table_headers = ["Status", "Action Label", "Seq", "Default(Y/N)", "Transition Condition", "Transition Role"]
            c = len(table_headers)
            count = rowcount(table)
            table.add_row().cells
            for i in range(c):
                hdr_cells = table.rows[count].cells
                run = hdr_cells[i].paragraphs[0].add_run(table_headers[i])
                run.bold = True
            for i in range(c):
                backgroundColor(table=table, row=count, column=i)
            nxtstat = value.get('nxtstatus')
            for nt in nxtstat:
                row_cells = table.add_row().cells
                row_cells[0].text = nt.get('nextStatus')
                row_cells[1].text = nt.get('boNextStatusLbl')
                row_cells[2].text = nt.get('sortSequence')
                row_cells[3].text = nt.get('shouldUseAsDefault')
                row_cells[4].text = nt.get('condition') if nt.get('condition') else ""
                row_cells[5].text = nt.get('role')

        count = rowcount(table)
        if 'statusoption' in value:
            table.add_row().cells
            hdr_cells = table.rows[count].cells
            run = hdr_cells[0].paragraphs[0].add_run('BO Status Option')
            run.bold = True
            row = table.row_cells(count)
            row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4]).merge(row[5])
            backgroundColor(table=table, row=count, column=0)

            table_headers = ["Option Type", "", "", "Seq", "", "Value"]
            c = len(table_headers)
            count = rowcount(table)
            table.add_row().cells
            for i in range(c):
                hdr_cells = table.rows[count].cells
                row = hdr_cells
                row[0].merge(row[1]).merge(row[2])
                row[4].merge(row[5])
                run = hdr_cells[i].paragraphs[0].add_run(table_headers[i])
                run.bold = True
            for i in range(c):
                backgroundColor(table=table, row=count, column=i)
            statusoption = value.get('statusoption')
            for statop in statusoption:
                row_cells = table.add_row().cells
                row = row_cells
                row[0].merge(row[1]).merge(row[2])
                row[4].merge(row[5])
                row_cells[0].text = statop.get('optionType')
                row_cells[3].text = statop.get('sequence')
                row_cells[4].text = statop.get('value')

        move_table_after(table, para)
        spaceAfterTable(table)


def setDataArea(para):
    p1 = tableHeader(para, "Data Area")

    for da in dataArea:
        table = create_table_header(da, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = dataArea[da][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Access Modes")
        value = dataArea[da][1]
        if value:
            row_cells = table1.add_row().cells
            row_cells[0].text = value
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)




def setUiMap(para):
    p1 = tableHeader(para, ComponentName.uiMap)

    table_headers = [ComponentName.uiMap, "Html", ComponentName.schema, ComponentName.messageDescription]
    table = create_table_header_and_row(row=1, column=4, style='Table Grid', list_of_headers=table_headers)
    table.autofit = True
    table.allow_autofit = True
    table.columns[0].width = ComponentName.columnWidth
    table.columns[1].width = 3000200
    for um in uiMap:
        data = uiMap[um]
        row_cells = table.add_row().cells
        row_cells[0].text = um
        row_cells[1].text = data[2]
        row_cells[2].text = data[0]
        row_cells[3].text = data[1]
    move_table_after(table, p1)
    spaceAfterTable(table)


def setMessage(para):
    p1 = tableHeader(para, ComponentName.message)
    table_headers = [ComponentName.messageCategory, ComponentName.messageNumber, ComponentName.messageDescription]
    table = create_table_header_and_row(row=1, column=3, style='Table Grid', list_of_headers=table_headers,
                                        list_of_rows=message_list)
    for i in range(len(table_headers)):
        backgroundColor(table=table, row=0, column=i)
    move_table_after(table, p1)
    spaceAfterTable(table)


def setCharType(para):
    p1 = tableHeader(para, ComponentName.characteristics)

    for char in charType_list:
        table = create_table_header(char, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = charType_list[char]
        if value:
            for entity in value:
                row_cells = table.add_row().cells
                row_cells[0].text = string.capwords(capital_words_spaces(entity))
                row_cells[1].text = value[entity]
        else:
            table.add_row().cells
        move_table_after(table, p1)
        spaceAfterTable(table)


def setScriptEntity(para):
    doc = Document(df.loc['templateDocument'].value)

    if 'OUAF Entities' in para.text:
        scriptStepskeys = list(scriptData.keys())
        script_steps=scriptData[scriptStepskeys[0]]
        script_steps_keys = list(script_steps.keys())
        list_of_info=['description','schema']
        script_description=script_steps['description']
        column = 0
        #uniqueScript = set({})
        p1 = tableHeader(para, ComponentName.scriptName)
        ScriptName = scriptStepskeys[0]

        para3 = insert_paragraph_after(p1, None)  # creating the paragraph
        para3.add_run(ScriptName).bold = True
        description='\n'+'Description:- '+script_description+'\n'
        para3.add_run(description)
        table = doc.add_table(rows=1, cols=2, style='Table Grid')  # create table
        table.cell(0, 0).width = ComponentName.columnWidth
        table.columns[0].width = ComponentName.columnWidth
        table.columns[1].width = 4800200
        hdr_cells = table.rows[0].cells
        hdr_cells[column].text = 'Steps'
        hdr_cells[column + 1].text = 'Script'

        for step_key in script_steps_keys:
            if step_key not in list_of_info:
                step_value = script_steps[step_key]
                row_cells = table.add_row().cells
                row_cells[column].text = step_key.split(ComponentName.seperator)[1]
                ##TODO- replace the code for this with code to pseudocode part
                steps = step_value
                row_cells[column + 1].text = steps
                move_table_after(table, para3)
                spaceAfterTable(table)


def setAlgo(para):
    paraValue = {}
    p1 = tableHeader(para, ComponentName.algorithm)

    for algo in algorithmSoftParam:
        table = create_table_header(algo, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = algorithmSoftParam[algo][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
            # Check if algorithm and alogorithm type name if same, If do so print parameter names to algorithm table
            if algorithmSoftParam[algo][1]:
                if entity == "algorithmType":
                    if algo == value[entity]:
                        algotyp = value[entity]
                        try:
                            paraValue = algoTypeSoftparam[algotyp][1]
                        except:
                            paraValue = {}
        move_table_after(table, p1)

        table1 = create_table_header("Parameters",
                                     ["Sequence", "Parameter Label", "value"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        algoTypePara = paraValue.get('algoTypPara')
        value = algorithmSoftParam[algo][1]
        values = value.get('algoPara')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                if algoTypePara:
                    for para in algoTypePara:
                        seq = para.get('sequence')
                        if seq == entity.get('sequence'):
                            row_cells[1].text = para.get('parameterLabel')
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[2].text = entity.get('value') if entity.get('value') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)




def setAlgoType(para):
    p1 = tableHeader(para, "Algorithm Types")

    for algoTyp in algoTypeSoftparam:
        table = create_table_header(algoTyp, ["Property", "Value"])
        row = table.row_cells(0)
        row[0].merge(row[1])

        value = algoTypeSoftparam[algoTyp][0]
        for entity in value:
            row_cells = table.add_row().cells
            row_cells[0].text = string.capwords(capital_words_spaces(entity))
            row_cells[1].text = value[entity]
        move_table_after(table, p1)

        table1 = create_table_header("Parameters",
                                     ["Sequence", "Parameter Label", "is Parameter Required"])
        row = table1.row_cells(0)
        row[0].merge(row[1]).merge(row[2])

        value = algoTypeSoftparam[algoTyp][1]
        values = value.get('algoTypPara')
        if values:
            for entity in values:
                row_cells = table1.add_row().cells
                row_cells[0].text = entity.get('sequence') if entity.get('sequence') else ""
                row_cells[1].text = entity.get('parameterLabel') if entity.get('parameterLabel') else ""
                row_cells[2].text = entity.get('isParameterRequired') if entity.get('isParameterRequired') else ""
        else:
            table1.add_row().cells
        move_after_table(table1, table)
        spaceAfterTable(table1)

    # p1 = componentHeader(para, ComponentName.algorithmType)
    # for algoType in algorithmType:
    #     key = algoType.split("\n")[0]
    #
    #     para1 = insert_paragraph_after(p1, None, None)
    #
    #     table = create_table_header(key, ["Property", "Value"])
    #     row = table.row_cells(0)
    #     row[0].merge(row[1])
    #     index=0
    #     for i in range(len(algoTypeSoftparam)):
    #         row_cells = table.add_row().cells
    #         algo_type_soft_param_key = list(algoTypeSoftparam.get(key)[index + 1].keys())[0]
    #         algo_type_soft_param_value = algoTypeSoftparam.get(key)[index + 1].get(algo_type_soft_param_key)
    #         index=index+1
    #         row_cells[0].text = algo_type_soft_param_key
    #         row_cells[1].text = algo_type_soft_param_value
    #
    #     move_table_after(table, para1)
    #     if list(algoTypeSoftparam.get(key)[0].keys()) != 'discription':
    #         table_headers = [ComponentName.sequenceNo, ComponentName.parameter, ComponentName.parameterType,
    #                          ComponentName.required]
    #         table1 = create_table_header_and_row(row=1, column=4, style='Table Grid', list_of_headers=table_headers)
    #         c = len(table_headers)
    #         for i in range(c):
    #             backgroundColor(table=table1, row=0, column=i)
    #
    #         algoTypeParam = algoTypeSoftparam[key][0]
    #         for param in algoTypeParam:
    #             row_cells = table1.add_row().cells
    #             row_cells[0].text = param
    #             row_cells[1].text = algoTypeParam[param][0]
    #             row_cells[2].text = "Customer Modification"
    #             row_cells[3].text = algoTypeParam[param][1]
    #
    #         move_after_table(table1, table)
    #         spaceAfterTable(table1)


def setSqlQuery(querys, para, doc):
    if 'SQLs' in para.text:
        p1 = para._p
        for query in reversed(querys):
            # add the heading of component and mark it as bold
            # add rest of the query
            for zoneSql in query:
                para1 = doc.add_paragraph(None)
                qr=zoneSql.split("\n")
                para1.add_run(ComponentName.zone + ComponentName.seperator + " " + qr[0] + "\n").bold = True
                para1.add_run(qr[1] + "\n")
                p1.addnext(para1._p)


def setExportableComp(para):
    if 'Exportable Components under the module' in para.text:
        table_headers = [ComponentName.sequenceNo, ComponentName.componentType, ComponentName.componentName]
        table = create_table_header_and_row(row=1, column=3, style='Table Grid', list_of_headers=table_headers,
                                            list_of_rows=componentList)
        # table.autofit = True
        # table.allow_autofit = True
        table.cell(0, 0).width = ComponentName.columnWidth
        table.columns[0].width = ComponentName.columnWidth
        # table.columns[2].width = 3090200
        for i in range(len(table_headers)):
            backgroundColor(table=table, row=0, column=i)
        move_table_after(table, para)


# This function print the XML code by indentation
def xmlPrettyPrint(text):
    element = ET.XML(text)
    ET.indent(element)
    prettyXml = ET.tostring(element, encoding='unicode', )
    return prettyXml


# This function insert the paragraph at certain insertion point
def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def create_table_header(mainHeader, subHeaderList=None):
    doc = Document(df.loc['templateDocument'].value)
    if subHeaderList is not None:
        row = len(subHeaderList)
        column = len(subHeaderList)
    else:
        row = 1
        column = 1
    table = doc.add_table(rows=row, cols=column, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run(mainHeader).bold = True
    backgroundColor(table=table, row=0, column=0)

    if subHeaderList is not None:
        c = len(subHeaderList)
        for i in range(c):
            hdr_cells = table.rows[1].cells
            hdr_cells[i].paragraphs[0].add_run(subHeaderList[i]).bold = True
            for i in range(c):
                backgroundColor(table=table, row=1, column=i)
    return table


# This function move the table after certain paragraph
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


# This function move the new table followed old table
def move_after_table(newTable, oldTable):
    ntbl, otbl = newTable._tbl, oldTable._tbl
    otbl.addnext(ntbl)


# This function count the table rows present in table
def rowcount(table):
    count = 0
    for row in table.rows:
        count = count + 1
    return count


# This function create the table structure on given row and columm count
def create_table_header_and_row(row, column, list_of_headers, style=None, list_of_rows=None, width=None):
    doc = Document(df.loc['templateDocument'].value)
    # create the header of the table automatically
    table = doc.add_table(rows=row, cols=column, style=style)
    # columnWidth(table)
    if width is not None:
        table.columns[1].width = width
    hdr_cells = table.rows[0].cells
    for i in range(column):
        # hdr_cells[i].text = list_of_headers[i]
        run = hdr_cells[i].paragraphs[0].add_run(list_of_headers[i])
        run.bold = True
        # run.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

    # create the row for the table in the doc
    if list_of_rows != None:
        for rows_data in list_of_rows:
            row_cells = table.add_row().cells
            for i in range(column):
                row_cells[i].text = rows_data[i]
    return table


# This function color the table cells background -- greenish:7BA79D
def backgroundColor(table, row, column):
    shading_elm = parse_xml(r'<w:shd {} w:fill="A8A8A8"/>'.format(nsdecls('w')))
    table.cell(row, column)._tc.get_or_add_tcPr().append(shading_elm)


def columnWidth(table, col, width):
    table.autofit = True
    table.allow_autofit = True
    table.columns[col].width = width
    # table.cell(0, 0).width = ComponentName.columnWidth
    # table.columns[0].width = ComponentName.columnWidth


def capital_words_spaces(str1):
    return re.sub(r"(\w)([A-Z])", r"\1 \2", str1)


def tableHeader(para, name):
    doc = Document(df.loc['templateDocument'].value)
    p1 = insert_paragraph_after(para, None, None)
    p1.style = doc.styles['Heading 2']
    run = p1.add_run(name + ComponentName.seperator + "\n")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(197, 125, 90)
    run.font.name = 'PT Sans Narrow'
    return p1


def spaceAfterTable(table):
    doc = Document(df.loc['templateDocument'].value)
    tbl = table._tbl
    p1 = doc.add_paragraph()
    tbl.addnext(p1._p)
